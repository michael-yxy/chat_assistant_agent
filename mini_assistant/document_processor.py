import pdfplumber
from docx import Document
from pathlib import Path
from typing import List, Dict
import logging
import io

logger = logging.getLogger(__name__)

try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
    logger.info("OCR support enabled (pytesseract and PIL available)")
except ImportError as e:
    OCR_AVAILABLE = False
    logger.warning(f"OCR not available: {e}")


class DocumentProcessor:
    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def load_document(self, file_path: Path) -> str:
        suffix = file_path.suffix.lower()

        if suffix == '.pdf':
            return self._load_pdf(file_path)
        elif suffix == '.docx':
            return self._load_docx(file_path)
        elif suffix == '.txt':
            return self._load_txt(file_path)
        elif suffix in ['.xlsx', '.xls']:
            return self._load_excel(file_path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

    def _load_pdf(self, file_path: Path) -> str:
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    # 首先尝试常规文本提取
                    page_text = page.extract_text()
                    
                    # 如果没有提取到文本，尝试OCR
                    if not page_text or page_text.strip() == "":
                        page_text = self._extract_text_with_ocr(page)
                    
                    if page_text and page_text.strip():
                        text += page_text + "\n"
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {e}")
        return text
    
    def _extract_text_with_ocr(self, page) -> str:
        """使用OCR从PDF页面提取文本"""
        if not OCR_AVAILABLE:
            logger.warning("OCR not available, cannot extract text from image-based PDF")
            return ""
        
        text = ""
        try:
            # 提取页面上的所有图像
            images = page.images
            
            if not images:
                return ""
            
            # 按垂直位置排序图像（从上到下）
            images.sort(key=lambda img: -img['top'])
            
            for img in images:
                try:
                    # 获取图像数据（使用pdfplumber的to_image方法）
                    img_obj = page.to_image(resolution=300).original
                    
                    # 处理CMYK颜色空间
                    if img_obj.mode == 'CMYK':
                        img_obj = img_obj.convert('RGB')
                    
                    # 使用pytesseract进行OCR
                    # 将图像转换为灰度以提高OCR准确性
                    img_gray = img_obj.convert('L')
                    img_text = pytesseract.image_to_string(img_gray, lang='chi_sim+eng')
                    
                    if img_text and img_text.strip():
                        text += img_text + "\n"
                        # 只处理第一个图像（通常是整个页面）
                        break
                except Exception as img_e:
                    logger.warning(f"Error processing image: {img_e}")
            
            logger.debug(f"OCR extracted {len(text)} characters from page")
        except Exception as e:
            logger.error(f"Error extracting text with OCR: {e}")
        
        return text

    def _load_docx(self, file_path: Path) -> str:
        text = ""
        try:
            doc = Document(file_path)
            # 处理段落
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            # 处理表格
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
                text += "\n"
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {e}")
        return text

    def _load_txt(self, file_path: Path) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error loading TXT {file_path}: {e}")
            return ""

    def _load_excel(self, file_path: Path) -> str:
        text = ""
        try:
            import pandas as pd
            df = pd.read_excel(file_path)
            text = df.to_string()
        except Exception as e:
            logger.error(f"Error loading Excel {file_path}: {e}")
        return text

    def split_text(self, text: str) -> List[str]:
        if not text or not text.strip():
            return []

        chunks = []
        start = 0
        text_len = len(text)

        while start < text_len:
            end = start + self.chunk_size
            chunk = text[start:end]

            if end < text_len:
                newline_pos = chunk.rfind('\n')
                if newline_pos != -1:
                    chunk = chunk[:newline_pos]
                    end = start + newline_pos + 1

            chunk = chunk.strip()
            if chunk:
                chunks.append(chunk)

            start = end - self.chunk_overlap
            if start < 0:
                start = 0

        return chunks

    def process_file(self, file_path) -> List[Dict]:
        # 确保 file_path 是 Path 对象
        if isinstance(file_path, str):
            file_path = Path(file_path)
        
        content = self.load_document(file_path)
        chunks = self.split_text(content)

        documents = []
        for idx, chunk in enumerate(chunks):
            documents.append({
                'content': chunk,
                'metadata': {
                    'source': file_path.name,
                    'chunk_id': idx
                }
            })

        return documents
