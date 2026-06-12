import pdfplumber
from docx import Document
from pathlib import Path
from typing import List, Dict, Union, Optional, Literal
import logging
import io
import re

logger = logging.getLogger(__name__)

try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
    logger.info("OCR support enabled (pytesseract and PIL available)")
except ImportError as e:
    OCR_AVAILABLE = False
    logger.warning(f"OCR not available: {e}")

try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
    logger.info("docx2txt support enabled for faster DOCX parsing")
except ImportError as e:
    DOCX2TXT_AVAILABLE = False
    logger.warning(f"docx2txt not available, falling back to python-docx: {e}")

try:
    import numpy as np
    NUMPY_AVAILABLE = True
except ImportError as e:
    NUMPY_AVAILABLE = False
    logger.warning(f"NumPy not available: {e}")

ChunkStrategy = Literal['fixed', 'hierarchical', 'semantic']


class DocumentProcessor:
    def __init__(
        self, 
        chunk_size: int = 500, 
        chunk_overlap: int = 50,
        strategy: ChunkStrategy = 'fixed',
        semantic_threshold: float = 0.7
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.strategy = strategy
        self.semantic_threshold = semantic_threshold

    def load_document(self, file_path: Path) -> str:
        suffix = file_path.suffix.lower()

        if suffix == '.pdf':
            return self._load_pdf(file_path)
        elif suffix == '.docx':
            return self._load_docx(file_path)
        elif suffix == '.txt':
            return self._load_txt(file_path)
        elif suffix in ['.xlsx', '.xls']:
            return self._load_excel(file_path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")
    
    def stream_chunks(self, file_path: Union[str, Path]):
        if isinstance(file_path, str):
            file_path = Path(file_path)
        
        suffix = file_path.suffix.lower()
        
        content = self.load_document(file_path)
        
        if self.strategy == 'fixed':
            yield from self._chunk_fixed(content)
        elif self.strategy == 'hierarchical':
            yield from self._chunk_hierarchical(file_path, content)
        elif self.strategy == 'semantic':
            yield from self._chunk_semantic(content)
        else:
            raise ValueError(f"Unknown chunk strategy: {self.strategy}")

    def _load_pdf(self, file_path: Path) -> str:
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    
                    if not page_text or page_text.strip() == "":
                        page_text = self._extract_text_with_ocr(page)
                    
                    if page_text and page_text.strip():
                        text += page_text + "\n"
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {e}")
        return text
    
    def _extract_text_with_ocr(self, page) -> str:
        if not OCR_AVAILABLE:
            logger.warning("OCR not available, cannot extract text from image-based PDF")
            return ""
        
        text = ""
        try:
            images = page.images
            
            if not images:
                return ""
            
            largest_img = max(images, key=lambda img: img.get('width', 0) * img.get('height', 0))
            
            img_obj = page.to_image(resolution=150).original
            
            if img_obj.mode != 'L':
                img_obj = img_obj.convert('L')
            
            img_text = pytesseract.image_to_string(img_obj, lang='chi_sim+eng')
            text = img_text.strip()
            
            logger.debug(f"OCR extracted {len(text)} characters from page")
        except Exception as e:
            logger.error(f"Error extracting text with OCR: {e}")
        
        return text

    def _load_docx(self, file_path: Path) -> str:
        try:
            doc = Document(file_path)
            
            paragraphs_text = [para.text + "\n" for para in doc.paragraphs if para.text.strip()]
            
            tables_text = []
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        table_rows.append(" | ".join(row_text))
                if table_rows:
                    tables_text.append("\n--- TABLE START ---\n" + "\n".join(table_rows) + "\n--- TABLE END ---\n\n")
            
            return "".join(paragraphs_text) + "".join(tables_text)
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {e}")
            return ""

    def _load_txt(self, file_path: Path) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error loading TXT {file_path}: {e}")
            return ""

    def _load_excel(self, file_path: Path) -> str:
        text = ""
        try:
            import pandas as pd
            df = pd.read_excel(file_path)
            text = df.to_string()
        except Exception as e:
            logger.error(f"Error loading Excel {file_path}: {e}")
        return text

    def _chunk_fixed(self, text: str) -> List[str]:
        if not text or not text.strip():
            return []

        chunks = []
        start = 0
        text_len = len(text)

        while start < text_len:
            end = start + self.chunk_size
            chunk = text[start:end]

            if end < text_len:
                newline_pos = chunk.rfind('\n')
                if newline_pos != -1:
                    chunk = chunk[:newline_pos]
                    end = start + newline_pos + 1

            chunk = chunk.strip()
            if chunk:
                chunks.append(chunk)

            start = end - self.chunk_overlap
            if start < 0:
                start = 0

        return chunks

    def _chunk_hierarchical(self, file_path: Path, text: str) -> List[str]:
        if not text or not text.strip():
            return []

        chunks = []
        sections = self._extract_sections(text)
        
        for section in sections:
            section_chunks = self._chunk_fixed(section)
            chunks.extend(section_chunks)
        
        return chunks

    def _extract_sections(self, text: str) -> List[str]:
        sections = []
        current_section = []
        
        lines = text.split('\n')
        
        section_pattern = re.compile(r'^#{1,3}\s+.+|^[一二三四五六七八九十]+[、.．]\s+.+|^[0-9]+[.．]\s+.+')
        
        for line in lines:
            if section_pattern.match(line):
                if current_section:
                    sections.append('\n'.join(current_section))
                    current_section = []
                current_section.append(line)
            else:
                current_section.append(line)
        
        if current_section:
            sections.append('\n'.join(current_section))
        
        return sections

    def _chunk_semantic(self, text: str) -> List[str]:
        if not text or not text.strip():
            return []

        sentences = self._split_into_sentences(text)
        
        if len(sentences) == 0:
            return self._chunk_fixed(text)

        chunks = []
        current_chunk = []
        current_length = 0
        
        for i, sentence in enumerate(sentences):
            sentence_len = len(sentence)
            
            if current_length + sentence_len <= self.chunk_size:
                current_chunk.append(sentence)
                current_length += sentence_len
            else:
                if current_chunk:
                    chunks.append(''.join(current_chunk))
                
                if sentence_len > self.chunk_size:
                    sub_chunks = self._split_long_sentence(sentence)
                    chunks.extend(sub_chunks)
                    current_chunk = []
                    current_length = 0
                else:
                    current_chunk = [sentence]
                    current_length = sentence_len
        
        if current_chunk:
            chunks.append(''.join(current_chunk))
        
        return chunks

    def _split_into_sentences(self, text: str) -> List[str]:
        sentence_endings = re.compile(r'([。！？\.!?])')
        parts = sentence_endings.split(text)
        sentences = []
        
        for i in range(0, len(parts) - 1, 2):
            sentence = parts[i] + parts[i + 1]
            sentence = sentence.strip()
            if sentence:
                sentences.append(sentence)
        
        if parts[-1].strip():
            sentences.append(parts[-1].strip())
        
        return sentences

    def _split_long_sentence(self, sentence: str) -> List[str]:
        chunks = []
        start = 0
        sentence_len = len(sentence)
        
        while start < sentence_len:
            end = start + self.chunk_size
            
            if end < sentence_len:
                comma_pos = sentence.rfind('，', start, end)
                if comma_pos != -1:
                    end = comma_pos + 1
                else:
                    semicolon_pos = sentence.rfind('；', start, end)
                    if semicolon_pos != -1:
                        end = semicolon_pos + 1
            
            chunk = sentence[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end
        
        return chunks

    def process_file(self, file_path: Union[str, Path]) -> List[Dict]:
        if isinstance(file_path, str):
            file_path = Path(file_path)
        
        content = self.load_document(file_path)
        
        if self.strategy == 'fixed':
            chunks = self._chunk_fixed(content)
        elif self.strategy == 'hierarchical':
            chunks = self._chunk_hierarchical(file_path, content)
        elif self.strategy == 'semantic':
            chunks = self._chunk_semantic(content)
        else:
            chunks = self._chunk_fixed(content)

        documents = []
        for idx, chunk in enumerate(chunks):
            documents.append({
                'content': chunk,
                'metadata': {
                    'source': file_path.name,
                    'chunk_id': idx,
                    'strategy': self.strategy
                }
            })

        return documents